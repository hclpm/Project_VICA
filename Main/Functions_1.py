import numpy as np
from PIL import Image
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import RGBColor
from openai import OpenAI
import keys
import datetime
import my_vidstream.streaming as streaming

exit = False

img = np.full((480,640,3), 0, np.int8)
img_pil = Image.fromarray(img.astype(np.uint8))

# global variances
recognized_text = ""
translated_text = ""

# charGPT용 text
dialogue_for_opponent = []
dialogue_for_self = []
dialogue_integrated = []
conversation_state = ""

summary = []
subject = ""
import numpy as np

# stop state
self_stop_state = ""
opponent_stop_state = ""


def subtitle(subject_received):
    print('subtitle activated')
    global recognized_text, translated_text, subject, exit
    subject = subject_received
    import speech_recognition as sr
    # initialize
    recognizer = sr.Recognizer()
    # microphone connect
    microphone = sr.Microphone()
    # speech recognition and translation into subtitle
    with microphone as source:
        while exit != True:
            print("Listening...")
            audio = recognizer.listen(source)
            try:
                # speech recognition
                recognized_text = recognizer.recognize_google(audio, language="ko-KR")
                print("You said:   ", recognized_text)
                # translation
                translated_text = PAPAGO_translator(recognized_text)
                print("In English: ", translated_text)
                dialogue_append()
            except sr.UnknownValueError:
                print("Failed to understand radio")
            
        print("translator_thread closed")

def PAPAGO_translator(message):
    import keys as my
    import urllib.request
    # import keys as my
    import json
    client_id = my.PAPAGO_ClientID()
    client_secret = my.PAPAGO_ClientSecret()
    encText = urllib.parse.quote(message)
    data = "source=ko&target=en&text=" + encText
    url = "https://openapi.naver.com/v1/papago/n2mt"
    request = urllib.request.Request(url)
    request.add_header("X-Naver-Client-Id",client_id)
    request.add_header("X-Naver-Client-Secret",client_secret)
    response = urllib.request.urlopen(request, data=data.encode("utf-8"))
    rescode = response.getcode()
    if(rescode==200):
        response_body = response.read()
        response_data = json.loads(response_body.decode('utf-8'))
        translated_text = response_data['message']['result']['translatedText']
        return translated_text
    else:
        print("Error Code:" + rescode)

###########################################################################

def giveEngtext():
    return translated_text
def giveKortext():
    return recognized_text

####################################################################################################

# 대화 기록 저장
def dialogue_append():
    global recognized_text, dialogue_for_self, dialogue_for_opponent, self_stop_state, opponent_stop_state
    if self_stop_state == "" or opponent_stop_state == "":
        time = datetime.datetime.now()
        dialogue_for_self.append(("ME:     " + recognized_text.rstrip("\n"), time))
        dialogue_for_opponent.append(("OPPONENT:" + recognized_text.rstrip("\n"), time))


def client_conversation_append(opponent_dialogue):
    global dialogue_for_self, dialogue_integrated
    for tuples in opponent_dialogue:
        dialogue_for_self.append(tuples)
    dialogue_integrated = sorted(dialogue_for_self, key= lambda x: x[1])
    dialogue_integrated = [item[0] for item in dialogue_integrated]
    print(f"dialogue_integrated: {dialogue_integrated}")
    ChatGPT_summarize()

def set_run_font(run, font_name, font_size, east_asia_font=None):
    run.font.name = font_name
    run.font.size = Pt(font_size)

    if east_asia_font:
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), east_asia_font)

def save_dialogue_to_docx(dialogue_integrated, document_path='PrivateFiles/Projects/Communication_Assistant/Main_13/ConversationDocs/dialogue_final.docx'):
    doc = Document()

    # 문서 제목 추가 (중앙 상단에 위치)
    title = doc.add_heading('VICA\n', level=1)
    title.alignment = 1  # 중앙 정렬

    # 대화 시작
    doc.add_heading('Dialogue', level=2)

    for line in dialogue_integrated:
        speaker, text = line.split(":", 1)
        paragraph = doc.add_paragraph()

        # speaker 파란색으로 변경
        run_speaker = paragraph.add_run(f"{speaker.strip()}: ")
        run_speaker.font.color.rgb = RGBColor(0, 150, 0)

        # text 색상 변경 및 줄바꿈 삽입
        run_text = paragraph.add_run(f"\n{text.strip()}")
        run_text.font.color.rgb = RGBColor(0, 0, 0)

        # OPPONENT의 대화인 경우 오른쪽 정렬 및 speaker와 text 위치 변경
        if speaker.strip().upper() == 'OPPONENT':
            paragraph.alignment = 2  # 오른쪽 정렬
            run_speaker.text = f"OPPONENT: "
            run_speaker.font.color.rgb = RGBColor(0, 150, 0)

    # 문서 저장
    doc.save(document_path)
    print(f"대화 내용이 '{document_path}' 파일로 저장되었습니다.")

def addSummary(summary, existing_document_path='PrivateFiles/Projects/Communication_Assistant/Main_13/ConversationDocs/dialogue_final.docx'):
    global exit
    try:
        # 기존 문서 열기
        doc = Document(existing_document_path)
    except FileNotFoundError:
        print("Unable to find dialogue file(.docx)")

    # 내용 추가
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(summary)

    # 문서 저장
    doc.save(existing_document_path)
    print("Summary is now included to your file")
    exit = True
    streaming.exit_true()


    

# ChatGPT에 요청
def ChatGPT_summarize():
    global dialogue_integrated, subject, summary, self_stop_state, recognized_text
    subject_str = "".join(subject)
    print(f"subject: {subject_str}")
    client = OpenAI(
        api_key=keys.ChatGPT_key(),
    )
    print(dialogue_integrated)
    request = ", ".join(dialogue_integrated)
    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": f"전체 내용의 주제의 대한 요약과 결론만 출력해줘.(주제 = {subject_str}, 대화내용: " + request + ")",
            }
        ],
        model="gpt-3.5-turbo",
    )
    summary = chat_completion.choices[0].message.content
    print(summary)
    save_dialogue_to_docx(dialogue_integrated)
    addSummary(summary)
    recognized_text = "99qq"
    self_stop_state = "stop"