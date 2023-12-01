from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls


def set_run_font(run, font_name, font_size, east_asia_font=None):
    run.font.name = font_name
    run.font.size = Pt(font_size)

    if east_asia_font:
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), east_asia_font)

def save_dialogue_to_docx(dialogue, output_path, font_name="맑은 고딕", font_size=10, east_asia_font=None):
    doc = Document()

    for line in dialogue:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(line)
        set_run_font(run, font_name, font_size, east_asia_font)

    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.save(output_path)
    print(f"대화 내용이 '{output_path}' 파일로 저장되었습니다.")

if __name__ == "__main__":
    dialogue = [
        "1. User1: 안녕하세요?",
        "2. User2: 네 반갑습니다.",
        "3. User1: 정의란 무엇이라고 생각하시나요?",
        "4. User2: 정의란 남에게 피해를 주지 않는 것입니다.",
        "5. User1: 제 생각은 좀 다릅니다.",
        "6. User1: 정의란 남에게 피해를 주지 않는 것이 아니라, 모든 사람이 자신이 행동한 것에 대한 책임을 지는 것이라고 생각합니다.",
        "7. User1: 모두의 입장과 가치관이 같을 수 없기 때문에, 남에게 피해를 입히는 것은 불가피합니다.",
        "8. User1: 하지만 그 행동에 대해 공평한 책임이 부여된다면, 비로소 정의로워진다고 생각합니다.",
        "9. User2: 저 역시 듣고보니 그 의견에 동의할 수 밖에 없군요. 인정합니다.",
        "10. User2: 정의란 모두에게 공정한 책임을 부여한다는 의견은 합리적인 것 같습니다."
    ]

    output_file_path = "C:/Users/js2-3/Desktop/Github/PrivateFiles/Projects/Communication_Assistant/Main_2/ConversationDocs/dialogue.docx"
    save_dialogue_to_docx(dialogue, output_file_path, font_name="맑은 고딕", font_size=10, east_asia_font='맑은 고딕')