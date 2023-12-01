from docx import Document

def add_summary_to_document(summary, existing_document_path='dialogue_document.docx'):
    try:
        # 기존 문서 열기
        doc = Document(existing_document_path)
    except FileNotFoundError:
        # 기존 문서가 없는 경우 새로운 문서 생성
        doc = Document()

    # 내용 추가
    doc.add_heading('대화 내용 요약', level=1)
    doc.add_paragraph(summary)

    # 문서 저장
    doc.save(existing_document_path)

if __name__ == "__main__":
    summary_text = """
    두 사람은 서로의 이름을 소개하고, 현재 러시아와 우크라이나의 전쟁에 대한 이야기를 나누었습니다.
    대화는 음식에 대한 주제로 전환되어, 상대방은 고구마를 튀겨서 먹는 것을 선호한다고 말했습니다.
    나는 감자튀김을 더 좋아한다고 했고, 상대방은 감자튀김이 몸에 좋지 않다고 충고했습니다.
    """

    add_summary_to_document(summary_text)