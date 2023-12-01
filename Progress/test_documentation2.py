from docx import Document
from docx.shared import RGBColor

def save_dialogue_to_docx(dialogue, document_path='dialogue_final.docx'):
    doc = Document()

    # 문서 제목 추가 (중앙 상단에 위치)
    title = doc.add_heading('회의록\n\n', level=1)
    title.alignment = 1  # 중앙 정렬

    # 대화 시작
    doc.add_heading('대화 내용', level=2)

    for line in dialogue:
        speaker, text = line.split(":", 1)
        paragraph = doc.add_paragraph()

        # speaker 파란색으로 변경
        run_speaker = paragraph.add_run(f"{speaker.strip()}: ")
        run_speaker.font.color.rgb = RGBColor(0, 150, 0)

        # text 색상 변경 및 줄바꿈 삽입
        run_text = paragraph.add_run(f"\n{text.strip()}")
        run_text.font.color.rgb = RGBColor(0, 0, 0)

        # OPPONENT의 대화인 경우 오른쪽 정렬 및 speaker와 text 위치 변경
        if speaker.strip().upper() == 'OPPONENT':
            paragraph.alignment = 2  # 오른쪽 정렬
            run_speaker.text = f"OPPONENT: "
            run_speaker.font.color.rgb = RGBColor(0, 150, 0)

    # 문서 저장
    doc.save(document_path)

if __name__ == "__main__":
    dialogue_data = [
        'OPPONENT: 나는 김밥을 좋아한다',
        'ME:     나는 김밥을 좋아한다',
        'OPPONENT: 하지만 오늘 김밥이 품절이었다',
        'ME:     하지만 오늘 김밥이 품절이었다',
        'OPPONENT: 그래서 오늘 라면을 먹을까',
        'ME:     그래서 오늘 라면을 먹을까',
        'OPPONENT: 라면은 맛이 없다',
        'ME:     라면은 맛이 없다',
        'OPPONENT: 라면 말고 사과를 먹어야겠다',
        'ME:     라면 말고 사과를 먹어야겠다'
    ]

    create_dialogue_document(dialogue_data)